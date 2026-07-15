"""
Generate a realistic DOCX test CV for Sprint 6.1 production verification.
Same role family as the PDF (Senior Backend Engineer) so the test exercises
the same kind of skills/match path. Uses python-docx (already installed).
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches


def make_cv() -> Document:
    doc = Document()

    # Set default font for the document.
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Heading 1: Name
    h = doc.add_heading('Marcus Chen', level=0)

    # Subtitle: current title
    p = doc.add_paragraph()
    run = p.add_run('Senior Backend Engineer')
    run.bold = True
    run.font.size = Pt(13)

    # Contact line
    doc.add_paragraph(
        'Email: marcus.chen@example.com   |   '
        'Phone: +1-206-555-0192   |   '
        'Seattle, WA   |   '
        'LinkedIn: linkedin.com/in/marcuschen'
    )

    # Summary
    doc.add_heading('Professional Summary', level=2)
    doc.add_paragraph(
        'Senior backend engineer with 9 years of experience designing and scaling '
        'distributed systems in Go and Python. Built and operated the order-processing '
        'pipeline at a large e-commerce company, owning throughput, reliability, and '
        'cost. Comfortable across the stack: Go, Python, PostgreSQL, Kafka, Kubernetes, '
        'gRPC. Mentors mid-level engineers; has driven architecture reviews for the last 4 years.'
    )

    # Work experience
    doc.add_heading('Work Experience', level=2)

    doc.add_paragraph().add_run('Senior Backend Engineer @ Amazon').bold = True
    doc.add_paragraph('Seattle, WA  |  March 2019 – Present')
    for bullet in [
        'Owned the order-fulfillment service (Go, gRPC, PostgreSQL); sustained 99.99% availability over 4 years across peak holiday traffic.',
        'Designed the multi-region active-active failover for the checkout path; cut p99 latency 38% during regional incidents.',
        'Led the migration from a monolithic Python service to 8 Go microservices behind a gRPC mesh; cut EC2 spend by $1.4M/year.',
        'Mentored 5 mid-level engineers; co-authored the team\'s RFC template and architecture-review checklist.',
    ]:
        doc.add_paragraph(bullet, style='List Bullet')

    doc.add_paragraph().add_run('Backend Engineer @ Stripe').bold = True
    doc.add_paragraph('South San Francisco, CA (Remote)  |  August 2016 – February 2019')
    for bullet in [
        'Built the API gateway tier for the Payments product (Go); handled 12k req/s sustained.',
        'Migrated the legacy Java billing service to Go; cut p99 latency from 480ms to 95ms.',
        'Wrote the team\'s observability playbook (Prometheus + Grafana); on-call MTTR dropped from 47min to 14min.',
    ]:
        doc.add_paragraph(bullet, style='List Bullet')

    doc.add_paragraph().add_run('Software Engineer @ Microsoft').bold = True
    doc.add_paragraph('Redmond, WA  |  July 2014 – July 2016')
    for bullet in [
        'Contributed to the Azure Storage client library (C#, .NET); shipped the retry-policy overhaul used by 50+ internal teams.',
        'Owned the C# bindings for the Storage REST API.',
    ]:
        doc.add_paragraph(bullet, style='List Bullet')

    # Education
    doc.add_heading('Education', level=2)
    doc.add_paragraph(
        'Bachelor of Science in Computer Science   |   '
        'University of Washington   |   2010 – 2014'
    )

    # Skills
    doc.add_heading('Skills', level=2)
    doc.add_paragraph(
        'Go, Python, gRPC, REST, PostgreSQL, Kafka, Redis, Kubernetes, Docker, '
        'AWS (EC2, RDS, S3, Lambda), Prometheus, Grafana, OpenTelemetry, '
        'Distributed Systems, Microservices, Architecture Review, Mentorship'
    )

    return doc


def main() -> int:
    out_path = Path('test-fixtures/cvs/marcus-chen-backend.docx')
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = make_cv()
    doc.save(str(out_path))
    size = out_path.stat().st_size
    print(f'wrote {out_path} ({size} bytes)')
    return 0


if __name__ == '__main__':
    sys.exit(main())
